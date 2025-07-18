import streamlit as st
from docx import Document
from io import BytesIO
import datetime
import requests
import re

# Configuração da interface
st.set_page_config(page_title="Gerador de Contratos", layout="centered")
st.title("📄 Gerador de Contratos")

# --- Funções utilitárias ---

# Remove pontos, traços e barras do CNPJ
def limpar_cnpj(cnpj):
    return re.sub(r'\D', '', cnpj)

# Faz a substituição das tags mesmo dentro de runs separados
def substituir_variaveis(doc, substituicoes):
    for p in doc.paragraphs:
        for chave, valor in substituicoes.items():
            if chave in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if chave in inline[i].text:
                        inline[i].text = inline[i].text.replace(chave, valor)

# Lê uma cláusula .docx de uma subpasta
def ler_clausula(nome_arquivo):
    try:
        doc = Document(f"Clausulas/{nome_arquivo}")
        return "\n".join([p.text for p in doc.paragraphs])
    except:
        return ""

# Consulta API da BrasilAPI
def buscar_dados_cnpj(cnpj):
    url = f"https://brasilapi.com.br/api/cnpj/v1/{cnpj}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()
            return {
                "razao_social": data['razao_social'],
                "logradouro": data['logradouro'],
                "numero": data['numero'],
                "complemento": data.get("complemento", ""),
                "bairro": data['bairro'],
                "cep": data['cep'],
                "cidade": data['municipio'],
                "uf": data['uf']
            }
        else:
            st.warning("CNPJ inválido ou não encontrado.")
            return None
    except Exception as e:
        st.error(f"Erro na API: {e}")
        return None

# --- Formulário ---
cnpj_input = st.text_input("CNPJ (com ou sem formatação)", max_chars=18)
cnpj = limpar_cnpj(cnpj_input)

dados = None
if len(cnpj) == 14:
    dados = buscar_dados_cnpj(cnpj)

if dados:
    razao_social = dados["razao_social"]
    endereco = f'{dados["logradouro"]}, {dados["numero"]}'
    complemento = dados["complemento"]
    cep = dados["cep"]
    cidade = dados["cidade"]
    uf = dados["uf"]

    with st.expander("📄 Dados da empresa carregados automaticamente", expanded=True):
        st.markdown(f"**Razão Social:** {razao_social}")
        st.markdown(f"**CNPJ:** {cnpj_input}")
        st.markdown(f"**Endereço:** {endereco}")
        st.markdown(f"**Complemento:** {complemento}")
        st.markdown(f"**CEP:** {cep}")
        st.markdown(f"**Cidade:** {cidade}")
        st.markdown(f"**UF:** {uf}")
else:
    razao_social = st.text_input("Razão Social")
    endereco = st.text_input("Endereço")
    complemento = st.text_input("Complemento")
    cep = st.text_input("CEP")
    cidade = st.text_input("Cidade")
    uf = st.text_input("UF")

executivo = st.text_input("Nome do Executivo")
honorario = st.text_input("Percentual de Honorário (ex: 10)")
data_contrato = st.date_input("Data do Contrato", value=datetime.date.today())
incluir_grossup = st.checkbox("Incluir cláusula Gross-up")

# --- Geração do documento ---
if st.button("Gerar Contrato"):
    try:
        doc = Document("Contratos/contrato_modelo.docx")

        substituicoes = {
            "[RAZAOSOCIAL]": razao_social,
            "[CNPJ]": cnpj,
            "[ENDEREÇO]": endereco,
            "[COMPLEMENTO]": complemento,
            "[CEP]": cep,
            "[CIDADE]": cidade,
            "[UF]": uf,
            "[HONORARIO]": honorario,
            "[EXECUTIVO]": executivo,
            "[DATA_CONTRATO]": data_contrato.strftime("%d/%m/%Y"),
            "[CLAUSULAGROSSUPP]": ler_clausula("clausula_grossup.docx") if incluir_grossup else ""
        }

        substituir_variaveis(doc, substituicoes)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Contrato gerado com sucesso!")
        st.download_button("📥 Baixar Contrato", data=buffer, file_name=f"Contrato_{cnpj}.docx")

    except Exception as e:
        st.error(f"❌ Erro ao gerar o contrato: {e}")
